#!/usr/bin/env python3
"""Genere les manuscrits .docx de test du formateur KDP.

    pip install python-docx
    python tools/make-fixtures.py

Les fichiers atterrissent dans app/kdp-formatter/__fixtures__/.
Ils sont volontairement petits mais couvrent les cas qui cassent le lecteur.
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = os.path.join(os.path.dirname(__file__), "..", "app", "kdp-formatter", "__fixtures__")

PROSE = (
    "Le vent tournait sur la lande, et Marceau comprit qu’il ne rentrerait pas avant la "
    "nuit. Il serra son col, compta ses pas, et laissa la maison disparaître derrière lui. "
    "Personne au village ne parlait plus de ce qui s’était passé l’hiver précédent, mais "
    "chacun savait exactement où se trouvait la barrière, et pourquoi elle restait fermée."
)

DIALOGUE = (
    "— Tu comptes vraiment y retourner ? demanda-t-elle sans lever les yeux.\n"
    "— Quelqu’un doit bien le faire, dit Marceau."
)


def base(title: str, author: str = "Antoine", lang: str = "fr-FR") -> Document:
    doc = Document()
    doc.core_properties.title = title
    doc.core_properties.author = author
    doc.core_properties.language = lang
    return doc


def add_chapter_body(doc: Document, paragraphs: int = 4) -> None:
    for i in range(paragraphs):
        doc.add_paragraph(PROSE)
        if i == 1:
            doc.add_paragraph(DIALOGUE)


def clean_novel() -> None:
    """Cas nominal : styles Titre 1, formatage inline, listes, citation."""
    doc = base("La Lande Grise")
    for c in ("Chapitre premier", "Chapitre deuxième", "Chapitre troisième"):
        doc.add_heading(c, 1)
        add_chapter_body(doc)
        doc.add_heading("Une scène", 2)
        p = doc.add_paragraph()
        p.add_run("En gras. ").bold = True
        p.add_run("En italique. ").italic = True
        p.add_run("Souligné.").underline = True
        doc.add_paragraph("Premier point", style="List Bullet")
        doc.add_paragraph("Deuxième point", style="List Bullet")
        doc.add_paragraph("Une citation qui respire.", style="Quote")
        sep = doc.add_paragraph("* * *")
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(PROSE)
    doc.save(os.path.join(OUT, "clean-novel.docx"))


def word_toc() -> None:
    """Sommaire tape a la main, avec pointilles et numeros de page."""
    doc = base("Le Sommaire Parasite")
    doc.add_heading("Table des matières", 1)
    for i in range(1, 4):
        doc.add_paragraph(f"Chapitre {i} ................... {i * 14}")
    for i in range(1, 4):
        doc.add_heading(f"Chapitre {i}", 1)
        add_chapter_body(doc, 2)
    doc.save(os.path.join(OUT, "word-toc.docx"))


def no_headings() -> None:
    """Aucun style Titre : le decoupage doit passer par l'heuristique."""
    doc = base("Sans Titres")
    for i in ("CHAPITRE UN", "CHAPITRE DEUX", "CHAPITRE TROIS"):
        doc.add_paragraph(i)
        add_chapter_body(doc, 3)
    doc.save(os.path.join(OUT, "no-headings.docx"))


def long_novel() -> None:
    """Assez long pour franchir un seuil du tableau de reliure KDP."""
    doc = base("Le Long Hiver")
    for c in range(1, 21):
        doc.add_heading(f"Chapitre {c}", 1)
        add_chapter_body(doc, 12)
    doc.save(os.path.join(OUT, "long-novel.docx"))


def typography() -> None:
    """Apostrophes droites, guillemets droits, points de suspension, tirets."""
    doc = base("Typographie Brute")
    doc.add_heading("Chapitre unique", 1)
    doc.add_paragraph(
        "Il repondit \"non\", puis se tut... Elle insista : \"Pourquoi ?\" "
        "Il n'avait rien a ajouter -- rien du tout ! Vraiment ?"
    )
    doc.add_paragraph(PROSE.replace("’", "'"))
    doc.save(os.path.join(OUT, "typography.docx"))


def illustrated() -> None:
    """Manuscrit avec images : une nette, une volontairement trop peu definie."""
    from docx.shared import Inches
    from PIL import Image

    tmp_big = os.path.join(OUT, "_tmp-big.png")
    tmp_small = os.path.join(OUT, "_tmp-small.png")
    # 2400 px de large : largement au-dessus des 300 DPI sur une colonne de 4,5 po.
    Image.new("RGB", (2400, 1500), (180, 200, 220)).save(tmp_big)
    # 240 px : impossible d'atteindre 300 DPI, doit declencher l'alerte.
    Image.new("RGB", (240, 150), (220, 180, 180)).save(tmp_small)

    doc = base("Livre Illustre")
    doc.add_heading("Chapitre avec images", 1)
    doc.add_paragraph(PROSE)
    doc.add_picture(tmp_big, width=Inches(4))
    doc.add_paragraph(PROSE)
    doc.add_picture(tmp_small, width=Inches(4))
    doc.add_paragraph(PROSE)
    doc.add_heading("Chapitre sans image", 1)
    add_chapter_body(doc, 2)
    doc.save(os.path.join(OUT, "illustrated.docx"))

    os.remove(tmp_big)
    os.remove(tmp_small)


def empty_doc() -> None:
    doc = base("Vide")
    doc.add_paragraph("")
    doc.save(os.path.join(OUT, "empty.docx"))


if __name__ == "__main__":
    os.makedirs(OUT, exist_ok=True)
    for fn in (clean_novel, word_toc, no_headings, long_novel, typography, illustrated, empty_doc):
        fn()
        print("ecrit :", fn.__name__)
